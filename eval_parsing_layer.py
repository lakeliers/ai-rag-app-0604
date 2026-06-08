import csv
import html
import json
from io import BytesIO, StringIO
from pathlib import Path

from docx import Document
from openpyxl import Workbook

import parsing_layer


REPORT_PATH = Path("reports/parsing_layer_report.html")


class MemoryUpload:
    def __init__(self, name, data, input_preview, parser_rule):
        self.name = name
        self._data = data
        self.input_preview = input_preview
        self.parser_rule = parser_rule

    def getvalue(self):
        return self._data


def make_markdown_upload():
    text = """# 产品 PRD

## 上线时间
产品预计 6 月 10 日上线。

## 价格策略
会员首月价格为 129 元，次月恢复原价。
"""
    return MemoryUpload(
        "product_prd.md",
        text.encode("utf-8"),
        input_preview=text,
        parser_rule="Markdown/TXT：按行读取，识别 #/## 标题和编号标题；标题之间的内容形成一个 ParsedSection。",
    )


def make_docx_upload():
    doc = Document()
    doc.add_heading("客服升级目标", level=1)
    doc.add_paragraph("客服升级后，首响时间控制在 30 秒以内。")
    doc.add_heading("工作台改版", level=1)
    doc.add_paragraph("工作台需要展示用户历史订单、会员等级和最近一次投诉。")
    buffer = BytesIO()
    doc.save(buffer)
    preview = """Word 文档结构：
Heading 1：客服升级目标
正文：客服升级后，首响时间控制在 30 秒以内。
Heading 1：工作台改版
正文：工作台需要展示用户历史订单、会员等级和最近一次投诉。"""
    return MemoryUpload(
        "customer_service_prd.docx",
        buffer.getvalue(),
        input_preview=preview,
        parser_rule="DOCX：读取段落 paragraph，根据 Word 样式 Heading/标题 判断小节标题；标题下正文形成 ParsedSection。",
    )


def make_csv_upload():
    buffer = StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["套餐", "首月价格", "原价"])
    writer.writerow(["基础版", "99", "199"])
    writer.writerow(["专业版", "129", "299"])
    writer.writerow(["企业版", "399", "599"])
    preview = buffer.getvalue()
    return MemoryUpload(
        "pricing.csv",
        preview.encode("utf-8"),
        input_preview=preview,
        parser_rule="CSV：读取行列数据，保留 row_start/row_end；当前样例行数较少，因此形成一个 table 类型 ParsedSection。",
    )


def make_xlsx_upload():
    workbook = Workbook()
    pricing = workbook.active
    pricing.title = "定价方案"
    pricing.append(["套餐", "首月价格", "原价"])
    pricing.append(["基础版", "99", "199"])
    pricing.append(["专业版", "129", "299"])

    launch = workbook.create_sheet("上线排期")
    launch.append(["模块", "上线时间"])
    launch.append(["会员页", "6 月 10 日"])
    launch.append(["支付页", "6 月 15 日"])

    buffer = BytesIO()
    workbook.save(buffer)
    preview = """Excel 工作簿结构：
Sheet：定价方案
套餐 | 首月价格 | 原价
基础版 | 99 | 199
专业版 | 129 | 299

Sheet：上线排期
模块 | 上线时间
会员页 | 6 月 10 日
支付页 | 6 月 15 日"""
    return MemoryUpload(
        "business_plan.xlsx",
        buffer.getvalue(),
        input_preview=preview,
        parser_rule="XLSX：先按 sheet 读取，再按行批次组织为 table 类型 ParsedSection，保留 sheet、row_start、row_end。",
    )


def make_json_upload():
    data = {
        "pricing": {
            "basic": 99,
            "pro": 129,
        },
        "launch": {
            "member_page": "6 月 10 日",
            "payment_page": "6 月 15 日",
        },
    }
    preview = json.dumps(data, ensure_ascii=False, indent=2)
    return MemoryUpload(
        "config.json",
        preview.encode("utf-8"),
        input_preview=preview,
        parser_rule="JSON：先 json.loads 解析结构；dict 按顶层 key 拆，list 按 item 拆。",
    )


def make_samples():
    return [
        make_markdown_upload(),
        make_docx_upload(),
        make_csv_upload(),
        make_xlsx_upload(),
        make_json_upload(),
    ]


def section_to_dict(section):
    return {
        "section_title": section.section_title,
        "content_type": section.content_type,
        "page": section.page,
        "sheet": section.sheet,
        "row_start": section.row_start,
        "row_end": section.row_end,
        "text": section.text,
    }


def evaluate_samples():
    results = []
    for upload in make_samples():
        sections = parsing_layer.read_upload_as_sections(upload)
        results.append({
            "file": upload.name,
            "input_preview": upload.input_preview,
            "parser_rule": upload.parser_rule,
            "section_count": len(sections),
            "sections": [section_to_dict(section) for section in sections],
        })
    return results


def render_metadata(section):
    metadata = {
        "section_title": section["section_title"],
        "content_type": section["content_type"],
        "page": section["page"],
        "sheet": section["sheet"],
        "row_start": section["row_start"],
        "row_end": section["row_end"],
    }
    rows = []
    for key, value in metadata.items():
        if value in ("", None):
            value = "空"
        rows.append(
            f"<tr><th>{html.escape(key)}</th><td>{html.escape(str(value))}</td></tr>"
        )
    return "<table class=\"metadata\"><tbody>" + "\n".join(rows) + "</tbody></table>"


def render_sections(sections):
    rows = []
    for index, section in enumerate(sections, start=1):
        rows.append(
            f"""
            <div class="section-output">
              <div class="section-title">ParsedSection {index}</div>
              <div class="section-grid">
                <div>
                  <div class="label">metadata</div>
                  {render_metadata(section)}
                </div>
                <div>
                  <div class="label">text</div>
                  <pre>{html.escape(section['text'])}</pre>
                </div>
              </div>
            </div>
            """
        )
    return "\n".join(rows)


def render_report(results):
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    cards = []
    for item in results:
        cards.append(
            f"""
            <section>
              <div class="section-head">
                <h2>{html.escape(item['file'])}</h2>
                <span>{item['section_count']} sections</span>
              </div>
              <div class="io-grid">
                <div class="panel input-panel">
                  <div class="label">原始输入预览</div>
                  <pre>{html.escape(item['input_preview'])}</pre>
                </div>
                <div class="panel">
                  <div class="label">解析规则</div>
                  <p>{html.escape(item['parser_rule'])}</p>
                </div>
              </div>
              <h3>解析输出</h3>
              {render_sections(item['sections'])}
            </section>
            """
        )

    html_text = f"""
<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <title>RAG 解析层评估报告</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; margin: 0; background: #f4f6f8; color: #1f2933; }}
    main {{ max-width: 1280px; margin: 0 auto; padding: 28px; }}
    .hero, section {{ background: #fff; border: 1px solid #dfe3e8; border-radius: 8px; padding: 22px; margin-bottom: 18px; }}
    h1, h2 {{ margin: 0; }}
    h3 {{ margin: 18px 0 10px; }}
    .section-head {{ display: flex; justify-content: space-between; gap: 16px; align-items: center; margin-bottom: 14px; }}
    .section-head span {{ background: #e3f2fd; color: #174ea6; border-radius: 999px; padding: 4px 10px; font-size: 13px; font-weight: 600; }}
    .io-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 14px; margin: 14px 0; }}
    .panel {{ border: 1px solid #dfe3e8; background: #fbfcfd; border-radius: 8px; padding: 14px; }}
    .input-panel {{ background: #fffdf5; }}
    .label {{ font-size: 12px; color: #66788a; font-weight: 700; text-transform: uppercase; margin-bottom: 8px; }}
    .section-output {{ border: 1px solid #dfe3e8; border-radius: 8px; margin-top: 12px; overflow: hidden; }}
    .section-title {{ background: #f5f7fa; padding: 10px 12px; font-weight: 700; }}
    .section-grid {{ display: grid; grid-template-columns: 300px 1fr; gap: 0; }}
    .section-grid > div {{ padding: 12px; }}
    .section-grid > div + div {{ border-left: 1px solid #dfe3e8; }}
    table {{ width: 100%; border-collapse: collapse; font-size: 13px; }}
    th, td {{ border: 1px solid #dfe3e8; padding: 8px; vertical-align: top; }}
    th {{ background: #f5f7fa; text-align: left; }}
    .metadata th {{ width: 120px; }}
    pre {{ margin: 0; white-space: pre-wrap; font-family: ui-monospace, SFMono-Regular, Menlo, monospace; line-height: 1.5; }}
    @media (max-width: 900px) {{ .io-grid, .section-grid {{ grid-template-columns: 1fr; }} .section-grid > div + div {{ border-left: 0; border-top: 1px solid #dfe3e8; }} }}
  </style>
</head>
<body>
  <main>
    <div class="hero">
      <h1>RAG 解析层评估报告</h1>
      <p>本报告只评估解析层：原始文件是否被转换成结构化 ParsedSection，不涉及切分、向量入库和模型回答。</p>
    </div>
    {''.join(cards)}
  </main>
</body>
</html>
"""
    REPORT_PATH.write_text(html_text, encoding="utf-8")


def main():
    results = evaluate_samples()
    render_report(results)
    print(json.dumps({
        "report": str(REPORT_PATH),
        "files": len(results),
        "summary": [
            {
                "file": item["file"],
                "section_count": item["section_count"],
                "titles": [section["section_title"] for section in item["sections"]],
            }
            for item in results
        ],
    }, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
